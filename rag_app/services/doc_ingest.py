"""Upload-and-condense service for user-supplied documents.

Flow:
    1. Browser uploads a .docx / .pdf / .md / .txt file to /session/attach.
    2. We parse it to plain text (``parse_bytes``).
    3. ``small_llm`` condenses the plain text into a dense markdown summary
       that preserves every number/observation but strips filler prose.
    4. Both raw + condensed are cached in an in-memory session store keyed
       by a generated ``session_id``; subsequent chat turns only ship
       ``session_id`` and the backend injects the condensed text into the
       evidence block. This keeps per-turn token cost bounded at ~500-800
       tokens instead of re-sending 3-5k tokens of document on every turn.

In-memory only: survives until the uvicorn process restarts. Fine for a
local POC. If/when we need persistence, swap the dict for a sqlite/redis
store with the same get/put/delete surface.
"""

from __future__ import annotations

import io
import re
import threading
import time
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, Optional

from langchain_core.messages import HumanMessage


# ---------------------------------------------------------------------------
# Parsing — dispatch by file extension / content-type
# ---------------------------------------------------------------------------

_DOCX_EXTS = {".docx"}
_PDF_EXTS = {".pdf"}
_TEXT_EXTS = {".md", ".markdown", ".txt"}
_ALLOWED_EXTS = _DOCX_EXTS | _PDF_EXTS | _TEXT_EXTS


def _ext(filename: str) -> str:
    m = re.search(r"\.[A-Za-z0-9]+$", filename or "")
    return m.group(0).lower() if m else ""


def parse_docx(data: bytes) -> str:
    """Extract paragraph text from a .docx byte blob."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables — flatten row-by-row so numeric comparison tables survive.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def parse_pdf(data: bytes) -> str:
    """Extract text from a PDF byte blob using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n".join(parts).strip()


def parse_text(data: bytes) -> str:
    try:
        return data.decode("utf-8").strip()
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore").strip()


def parse_bytes(filename: str, data: bytes) -> str:
    """Dispatch to the right parser. Raises ValueError for unknown ext."""
    ext = _ext(filename)
    if ext in _DOCX_EXTS:
        return parse_docx(data)
    if ext in _PDF_EXTS:
        return parse_pdf(data)
    if ext in _TEXT_EXTS:
        return parse_text(data)
    raise ValueError(
        f"Unsupported file extension '{ext}'. Allowed: {sorted(_ALLOWED_EXTS)}"
    )


# ---------------------------------------------------------------------------
# Condensation — one small_llm call, preserves numbers, drops filler
# ---------------------------------------------------------------------------

_CONDENSE_PROMPT = """\
You will compress the user's experimental document into a dense structured
markdown summary that another LLM will reference in subsequent Q&A turns.
The goal is RETENTION, not summarization — another LLM must be able to
answer questions about this experiment using only your output.

MUST PRESERVE (every instance, verbatim where possible):
- All numerical values (cell density, volumes, concentrations, incubation
  times, temperatures, pH, fold-changes, OD/OCR/ECAR readings, etc.)
- All experimental condition comparisons (Exp A / B / C style groupings) —
  use a markdown table when ≥3 conditions are compared.
- All observations, anomalies, unexpected results, and user questions
  embedded in the document.
- All reagent identities, equipment, plate formats, lot numbers.
- **The user's own interpretation, analysis, conclusions, and any named
  frameworks or categorizations they introduce** — e.g. if the user labels
  experimental regimes ("suppressed dynamic range", "most reliable system",
  "problematic system"), classifies conditions into categories, or draws
  cross-experiment synthesis, you MUST preserve those labels and
  categorizations verbatim. These are the user's intellectual framing and
  the downstream LLM will be expected to ADOPT them, not re-derive.
- Any section the user titles "Interpretation", "Synthesis", "Summary",
  "Conclusion", "Cross-experiment", or similar — copy these through
  (lightly compressed) rather than dropping them as "filler".

DROP:
- Introductions, background science the user already knows.
- Repeated descriptions of the same fact.
- Politeness / transition sentences.

NEVER DROP:
- A conclusion or interpretation section, even if short.
- Named categories / labels the user invented.

OUTPUT FORMAT:
- Start with one line: "Document: <inferred title or filename>"
- "## Setup" — bullet list of key parameters.
- "## Conditions" — markdown table if multiple conditions.
- "## Observations" — bullet list of every number + observation.
- "## User's interpretation" — preserve the user's own analysis,
  categorization, and conclusions verbatim or lightly compressed. If the
  user named regimes/categories, list them here with their definitions.
- "## Open questions" — anything the user is asking or flagging as
  uncertain.

Density over prose. Answer in the same language as the input document.

Document:
---
{raw_text}
---
"""


def condense_document(raw_text: str, small_llm: Any, *, max_chars: int = 24000) -> str:
    """Run one small_llm call to compress ``raw_text`` into a dense summary.

    ``max_chars`` caps the raw input sent to the LLM so absurdly long
    uploads don't blow the small-model context. Overflow is truncated with
    a clear marker so downstream users can see content was cut.
    """
    if not raw_text or not raw_text.strip():
        return ""
    if small_llm is None:
        # Degrade gracefully: return the raw text clipped. Better than
        # failing the upload outright during local dev.
        return raw_text[:max_chars]

    truncated = raw_text
    if len(raw_text) > max_chars:
        truncated = raw_text[:max_chars] + "\n\n[... truncated ...]"

    prompt = _CONDENSE_PROMPT.format(raw_text=truncated)
    try:
        resp = small_llm.invoke([HumanMessage(content=prompt)])
        out = resp.content if hasattr(resp, "content") else str(resp)
        return (out or "").strip()
    except Exception:
        # If the condense call fails, fall back to truncated raw text so
        # the upload is still useful.
        return truncated


# ---------------------------------------------------------------------------
# Session store — {session_id: DocSession}, in-memory, thread-safe
# ---------------------------------------------------------------------------


@dataclass
class DocSession:
    session_id: str
    filename: str
    raw: str
    condensed: str
    char_count: int
    created_at: float = field(default_factory=time.time)

    def to_public(self) -> Dict[str, Any]:
        """Trim to the fields that are safe/useful to return to the client."""
        # Return a short preview of the condensed doc so the user can sanity
        # check what got uploaded without paying for the full text.
        preview = self.condensed[:600]
        if len(self.condensed) > 600:
            preview += " …"
        # Rough token estimate (English ≈ 4 chars/token; we're conservative).
        token_estimate = max(1, len(self.condensed) // 4)
        return {
            "session_id": self.session_id,
            "filename": self.filename,
            "char_count": self.char_count,
            "condensed_chars": len(self.condensed),
            "condensed_preview": preview,
            "token_estimate": token_estimate,
        }


class _SessionStore:
    def __init__(self) -> None:
        self._lock = threading.Lock()
        self._data: Dict[str, DocSession] = {}

    def put(self, session: DocSession) -> None:
        with self._lock:
            self._data[session.session_id] = session

    def get(self, session_id: str) -> Optional[DocSession]:
        if not session_id:
            return None
        with self._lock:
            return self._data.get(session_id)

    def delete(self, session_id: str) -> bool:
        with self._lock:
            return self._data.pop(session_id, None) is not None


# Module-level singleton — same process-wide store across requests.
_STORE = _SessionStore()


def get_session(session_id: str) -> Optional[DocSession]:
    return _STORE.get(session_id)


def delete_session(session_id: str) -> bool:
    return _STORE.delete(session_id)


def create_session(filename: str, raw: str, condensed: str) -> DocSession:
    sess = DocSession(
        session_id=uuid.uuid4().hex,
        filename=filename,
        raw=raw,
        condensed=condensed,
        char_count=len(raw),
    )
    _STORE.put(sess)
    return sess


# ---------------------------------------------------------------------------
# Evidence injection helper — used by the streaming chat route
# ---------------------------------------------------------------------------


def render_user_doc_block(session: DocSession) -> str:
    """Format a DocSession as the top section of the evidence block.

    Includes an inline instruction block immediately after the document so
    the graph LLM can't miss it — SYSTEM_PROMPT instructions get lost when
    evidence is long (PubMed abstracts + protocols can push past 6k tokens
    before the doc even arrives). Reminding the model right next to the
    content it's supposed to anchor to works better in practice.
    """
    return (
        "## User-provided experimental context (HIGHEST PRIORITY)\n"
        f"(Uploaded document: {session.filename})\n\n"
        f"{session.condensed}\n\n"
        "### How to use the document above\n"
        "- The document is the user's own experimental write-up, including "
        "their own interpretation and analysis.\n"
        "- If the document introduces named categories, regimes, or labels "
        "(e.g. 'suppressed dynamic range', 'suppressed ceiling', 'most reliable "
        "system', 'three regimes', 'ceiling effect', 'problematic system'), "
        "you MUST reuse those exact phrases verbatim in your answer. These are "
        "the user's own framing, not yours to rewrite.\n"
        "- Do NOT re-derive conclusions the user has already drawn. Treat their "
        "'Interpretation' / 'Synthesis' / 'Conclusion' sections as premises.\n"
        "- Open your answer by classifying the user's situation using their "
        "own labels (e.g. 'Your Exp C falls in the *suppressed dynamic range* "
        "regime you identified…'), then build the specific answer on top.\n"
        "- Numbers from this document (cell density, OCR, timing) take "
        "precedence over anything in the PubMed/protocol sections below.\n"
    )
